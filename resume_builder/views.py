from django.shortcuts import render, redirect
from users.models import UserProfile
from django.contrib.auth.decorators import login_required
from .utils import enhance_with_ollama
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
import io
from .models import resume

@login_required
def generate_resume(request):
    try:
        profile = request.user.userprofile
    except UserProfile.DoesNotExist:
        return redirect('edit_profile')

    template_choice = profile.resume_template_choice or 1
    template_name = f'resume_templates/template{template_choice}.html'

    # Enhance summary (optional)
    enhanced_projects = enhance_with_ollama(
        f"{profile.projects}"
    )
    enhanced_achievements = enhance_with_ollama(
        f"Rewrite the following achievements in a concise and impactful way:\n{profile.achievements}"
    )

    context = {
        'full_name': profile.full_name,
        'phone_number': profile.phone_number,
        'email': request.user.email,
        'location': profile.location,
        'skills': profile.skills,
        'education': profile.education,
        'experience': profile.experience,
        'projects': enhanced_projects,
        'achievements': enhanced_achievements or profile.achievements,
    }
    resume.objects.create(
    user=request.user,
    full_name=profile.full_name,
    phone_number=profile.phone_number,
    email=request.user.email,
    location=profile.location,
    skills=profile.skills,
    education=profile.education,
    experience=profile.experience,
    projects=enhanced_projects,
    achievements=enhanced_achievements or profile.achievements
)
    
    
    

    return render(request, template_name, context)

from django.http import HttpResponse
from django.template.loader import get_template
# from weasyprint import HTML

# @login_required
# def download_resume_pdf(request):
#     try:
#         profile = request.user.userprofile
#     except UserProfile.DoesNotExist:
#         return redirect('edit_profile')

#     template_choice = profile.resume_template_choice or 1
#     template_path = f'resume_templates/template{template_choice}.html'

#     enhanced_projects = enhance_with_ollama(
#         f"{profile.projects}"
#     )
#     enhanced_achievements = enhance_with_ollama(
#         f"{profile.achievements}"
#     )

#     context = {
#         'full_name': profile.full_name,
#         'phone_number': profile.phone_number,
#         'email': request.user.email,
#         'location': profile.location,
#         'skills': profile.skills,
#         'education': profile.education,
#         'experience': profile.experience,
#         'projects': enhanced_projects or profile.projects,
#         'achievements': enhanced_achievements or profile.achievements,
#     }

#     html_template = get_template(template_path)
#     html_content = html_template.render(context)
#     pdf_file = HTML(string=html_content).write_pdf()

#     response = HttpResponse(pdf_file, content_type='application/pdf')
#     response['Content-Disposition'] = 'attachment; filename="resume.pdf"'
#     return response

import io
import pypandoc
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.contrib.auth.decorators import login_required
from django.shortcuts import redirect
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.shortcuts import redirect
import pypandoc
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from docx import Document
from io import BytesIO



@login_required
def download_resume_docx(request):
    try:
        profile = request.user.userprofile
    except UserProfile.DoesNotExist:
        return redirect('edit_profile')

    from .utils import enhance_with_ollama
    enhanced_projects = enhance_with_ollama(profile.projects)
    enhanced_achievements = enhance_with_ollama(profile.achievements)

    # Create DOCX file in memory
    document = Document()

    document.add_heading(profile.full_name, 0)
    document.add_paragraph(f"Email: {request.user.email}")
    document.add_paragraph(f"Phone: {profile.phone_number}")
    document.add_paragraph(f"Location: {profile.location}")
    document.add_paragraph(" ")

    # Skills
    document.add_heading('Skills', level=1)
    document.add_paragraph(profile.skills or "Not provided")

    # Education
    document.add_heading('Education', level=1)
    document.add_paragraph(profile.education or "Not provided")

    # Experience
    document.add_heading('Experience', level=1)
    document.add_paragraph(profile.experience or "Not provided")

    # Projects
    document.add_heading('Projects', level=1)
    document.add_paragraph(enhanced_projects or "Not provided")

    # Achievements
    document.add_heading('Achievements', level=1)
    document.add_paragraph(enhanced_achievements or "Not provided")

    # Prepare in-memory file
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Serve as download
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="resume.docx"'
    return response
# @login_required
# def download_resume_pdf(request):
#     try:
#         profile = request.user.userprofile
#     except UserProfile.DoesNotExist:
#         return redirect('edit_profile')

#     template_choice = profile.resume_template_choice or 1
#     template_path = f'resume_templates/template{template_choice}.html'

#     # Enhanced content (or fallback)
#     from .utils import enhance_with_ollama
#     enhanced_projects = enhance_with_ollama(profile.projects)
#     enhanced_achievements = enhance_with_ollama(profile.achievements)

#     context = {
#         'full_name': profile.full_name,
#         'phone_number': profile.phone_number,
#         'email': request.user.email,
#         'location': profile.location,
#         'skills': profile.skills,
#         'education': profile.education,
#         'experience': profile.experience,
#         'projects': enhanced_projects or profile.projects,
#         'achievements': enhanced_achievements or profile.achievements,
#     }

#     template = get_template(template_path)
#     html = template.render(context)

#     response = HttpResponse(content_type='application/pdf')
#     response['Content-Disposition'] = 'attachment; filename="resume.pdf"'

#     pisa_status = pisa.CreatePDF(io.StringIO(html), dest=response)
#     if pisa_status.err:
#         return HttpResponse('Error generating PDF', status=500)
#     return response


def view_resume(request):
    allresume = resume.objects.filter(user=request.user)
    return render(request, 'resume_templates/view_resume.html', {'allresume': allresume})

def resume_details(request, pk):
    try:
        resume1 = resume.objects.get(pk=pk, user=request.user)
    except resume.DoesNotExist:
        return redirect('view')

    context = {
        'full_name': resume1.full_name,
        'phone_number': resume1.phone_number,
        'email': resume1.email,
        'location': resume1.location,
        'skills': resume1.skills,
        'education': resume1.education,
        'experience': resume1.experience,
        'projects': resume1.projects,
        'achievements':resume1.achievements,
    }
    myuser = request.user
    myprofile = UserProfile.objects.get(user=myuser)
    
    template = myprofile.resume_template_choice
    return render(request, f'resume_templates/template{template}.html', context)


def delete_resume(request, pk):
    try:
        resume1 = resume.objects.get(pk=pk, user=request.user)
        resume1.delete()
    except resume.DoesNotExist:
        return redirect('view')

    return redirect('view')